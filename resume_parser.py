"""Text extraction helpers for uploaded resume files (PDF / DOCX / TXT)."""

import io


def extract_text(uploaded_file) -> str:
    """
    Extract plain text from a Streamlit UploadedFile.
    Supports .pdf, .docx, and .txt based on the file's extension.
    """
    name = (uploaded_file.name or "").lower()
    data = uploaded_file.getvalue()

    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore")

    raise ValueError(f"Unsupported file type: {name}")


def _extract_pdf(data: bytes) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError(
            "No extractable text found in this PDF (it may be a scanned image). "
            "Try pasting the resume text manually instead."
        )
    return text


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs]

    # Also pull text out of tables, since resumes often use them for layout.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)

    text = "\n".join(p for p in paragraphs if p.strip())
    if not text:
        raise ValueError("No extractable text found in this DOCX file.")
    return text
